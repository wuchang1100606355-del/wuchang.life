#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Generate a local patent filing draft package for review.

The generator intentionally takes applicant identity data from environment
variables so the source file does not store personal identifiers.
"""

from __future__ import annotations

import hashlib
import json
import os
from datetime import datetime, timezone, timedelta
from pathlib import Path
from textwrap import dedent

from docx import Document
from docx.shared import Pt


TZ = timezone(timedelta(hours=8))
NOW = datetime.now(TZ)
ROOT = Path("/home/taiji_admin/Taiji_Hub")
OUT = ROOT / "patent_filing" / "readonly_boot_edge_runtime_m5_v0_1"


TITLE = "一種基於唯讀啟動媒體與五維度規封包之邊緣運算執行環境注入、瞬態 I/O 記憶體調度及可稽核狀態回滾方法"
SHORT_TITLE = "唯讀啟動媒體五維度規邊緣執行方法"


def env(name: str, default: str = "") -> str:
    return os.environ.get(name, default).strip()


INVENTOR = {
    "name": env("TAIJI_PATENT_INVENTOR_NAME", "【發明人姓名】"),
    "id": env("TAIJI_PATENT_INVENTOR_ID", "【身分證字號】"),
    "address": env("TAIJI_PATENT_INVENTOR_ADDRESS", "【發明人地址】"),
}

BUSINESS = {
    "name": env("TAIJI_PATENT_BUSINESS_NAME", "【商號名稱】"),
    "ubn": env("TAIJI_PATENT_BUSINESS_UBN", "【統一編號】"),
    "address": env("TAIJI_PATENT_BUSINESS_ADDRESS", "【商號地址】"),
    "representative": env("TAIJI_PATENT_BUSINESS_REPRESENTATIVE", INVENTOR["name"]),
}


ABSTRACT = dedent(
    f"""
    本發明係關於一種邊緣運算執行環境建置與治理方法，係利用唯讀啟動媒體啟動邊緣運算裝置，擷取使用者資料、全域環境設定及系統預設參數，並將該等參數轉換為包含節點或身份、資料敏感性、動作意圖、權限時間窗口及可逆性或公共價值之五維度規封包。系統依該五維度規封包注入對應執行環境，並於揮發性記憶體中建立瞬態輸入輸出區，以供應用程式、容器、邊緣節點或人工智慧執行期暫存讀寫。執行期間，邊緣 I/O 記憶體調度器依據記憶體壓力、讀寫壓力、快取命中率、模型載入成本、資料敏感性及重放風險，決定本地執行、雲端轉送、延後執行、阻斷或隔離。稽核記錄模組記錄封包雜湊、父雜湊、啟動工作階段、重放防護參數及回滾參照。於終止條件成立時，回復模組清除瞬態輸入輸出區、停止相關服務、封存必要稽核摘要並回復系統狀態。藉此，可於資源受限之邊緣節點中提供可重複部署、低持久寫入、無明文上下文外洩風險且具可稽核回滾能力之執行環境。
    """
).strip()


SPECIFICATION = dedent(
    f"""
    【發明名稱】
    {TITLE}

    【技術領域】
    本發明係有關一種邊緣運算執行環境建置、人工智慧執行期治理、記憶體輸入輸出調度與狀態回復技術，尤指一種利用唯讀啟動媒體進行環境參數擷取、五維度規封包產生、執行環境注入、瞬態 I/O 記憶體調度、稽核記錄及執行狀態回滾之方法。

    【先前技術】
    現有邊緣運算節點多仰賴可寫入之本地儲存媒體或持久性檔案系統以部署作業系統、應用程式及相關設定檔，並藉由編修設定檔、系統環境變數或容器映像檔以完成執行環境建置。然此類方式通常需事先完成安裝、設定與後續維護，於節點數量增加、環境需頻繁更新或分散式部署時，易造成部署流程複雜、管理負擔提高及設定不一致之情形。

    此外，持久性儲存媒體於運作期間若遭受非預期寫入、惡意竄改或異常損毀，可能影響系統穩定性與安全性。既有容器化部署方案雖可支援部分環境變數注入機制，惟多仍依賴可寫入基礎系統或持久性映像檔，且對於資源受限之邊緣裝置，未必能兼顧瞬態暫存需求、執行後狀態清理、人工智慧模型載入成本、重放防護及環境一致性維持。

    又，現有人工智慧服務多以自然語言、明文上下文或雲端推理請求作為主要交換資料，當該等服務應用於社區、物業、商業或其他邊緣治理場域時，若未能於執行前以資料敏感性、權限窗口、回滾性及拓樸位置進行統一判斷，可能產生明文資料外洩、未授權執行、狀態重放或不可追溯之風險。是以，如何提供一種可藉由唯讀啟動媒體完成執行環境建置，並以可計算之度規封包管理執行、記憶體 I/O、稽核及回滾者，實為相關技術領域所待解決之課題。

    【發明內容】
    本發明之主要目的，在於提供一種基於唯讀啟動媒體與五維度規封包之邊緣運算執行環境建置方法，以降低部署複雜度，減少持久性寫入所衍生之風險，提升多次啟動時之環境一致性，並使人工智慧或邊緣服務之執行具備資料最小化、稽核追蹤、重放防護及狀態回復能力。

    為達上述目的，本發明係於邊緣運算裝置上以唯讀啟動媒體進行啟動，並由環境參數擷取模組自使用者資料、全域環境設定及系統預設設定中取得執行參數。封包產生模組將該等參數及待執行動作轉換為五維度規封包 M5，其中 M5 至少包含：節點或身份 N、資料敏感性 D、動作意圖 A、權限時間窗口 W 及可逆性或公共價值 R。AI 匝道判決模組依據 M5 及政策基準產生治理判決 Vr，該治理判決至少包含允許、稽核、警示、阻斷、要求人類確認或要求回滾之其中一者。

    執行環境注入模組依據該五維度規封包及治理判決建立或調整執行環境。於一實施態樣中，執行環境注入模組可將相關參數注入容器協調檔案所定義之服務，以啟動對應之應用程式、容器或人工智慧執行期。於執行期間，系統於揮發性記憶體中建立瞬態輸入輸出區，以供暫存資料、模型執行中檔案、快取、日誌或中間結果使用，避免寫入唯讀啟動媒體或未授權持久性儲存區。

    邊緣 I/O 記憶體調度器依據治理判決 Vr、記憶體壓力、讀寫壓力、快取命中率、模型載入成本、網路成本、資料敏感性及重放風險，控制是否允許排載記憶體區段、是否允許讀取上下文片段、是否允許重組明文上下文、是否要求本地人工智慧執行期取得最小必要上下文，或是否將事件導入隔離狀態。稽核記錄模組記錄五維度規封包、治理判決、政策版本、節點識別、時間戳、調度結果、回滾狀態、封包雜湊、父雜湊、啟動工作階段識別碼及重放防護參數。當應用程式執行完成、異常中止或接獲終止指令時，回復模組關閉相關服務、清除瞬態輸入輸出區、移除執行期間注入之暫時性設定，並使系統狀態回復至預定之初始狀態。若偵測權限不連續、拓樸不一致、重放風險或資料敏感性不符，則將該執行事件封存至隔離狀態而非直接恢復執行。

    藉此，本發明可於每次啟動時提供一致之執行環境，維持唯讀啟動媒體不受非預期寫入而改變，並透過五維度規封包使邊緣人工智慧、容器服務與資料流之執行具備統一治理、無明文上下文外洩防護、可稽核及可回滾之技術效果。

    【圖式簡單說明】
    圖1係本發明之系統架構方塊圖。
    圖2係本發明之五維度規向量封包資料結構示意圖。
    圖3係本發明之 AI 匝道判決流程圖。
    圖4係本發明之邊緣 I/O 記憶體調度流程圖。
    圖5係本發明之稽核紀錄與回滾流程示意圖。

    【實施方式】
    以下配合圖式說明本發明之一較佳實施方式，惟本發明並不以此為限。如圖1至圖5所示，本發明之方法係先提供一唯讀啟動媒體100，該唯讀啟動媒體100內可包含啟動程式、必要之系統元件、預設設定資訊、政策基準130及容器協調檔案190。當邊緣運算裝置藉由該唯讀啟動媒體100啟動後，環境參數擷取模組110自使用者資料170、全域環境設定180及系統預設設定中讀取並整合執行所需參數。

    封包產生模組120依據前述參數及待執行動作產生五維度規封包 M5。該五維度規封包 M5 至少包含 N、D、A、W、R 五個欄位，其中 N 表示節點或身份，D 表示資料敏感性，A 表示動作意圖，W 表示權限時間窗口，R 表示可逆性或公共價值。於另一實施態樣中，五維度規封包 M5 可進一步包含政策雜湊值、節點簽章、封包有效期限、稽核識別碼及重放防護參數。

    AI 匝道判決模組140讀取五維度規封包 M5 及政策基準130，計算治理判決 Vr。該治理判決 Vr 可包含允許、稽核、警示、阻斷、要求人類確認及要求回滾等欄位。舉例而言，當 D 顯示資料敏感性高且 W 顯示權限時間窗口已逾期時，治理判決 Vr 可為阻斷；當 R 顯示不可逆且涉及公共風險時，治理判決 Vr 可要求人類確認；當 D 允許但仍需追蹤時，治理判決 Vr 可要求建立稽核紀錄。

    邊緣 I/O 記憶體調度器150依據治理判決 Vr 控制瞬態輸入輸出區130之配置、讀寫權限及生命週期。於一實施態樣中，邊緣 I/O 記憶體調度器150可取得記憶體壓力、讀取壓力、寫入壓力、快取命中率、模型載入成本、網路成本、磁碟成本、重放風險及稽核需求，以決定是否允許排載記憶體區段、是否讀取上下文片段、是否重組明文上下文、是否啟動人工智慧執行期160、是否將任務導向邊緣節點或儲存節點190，或是否將事件移入隔離狀態。

    執行環境140得包括容器、服務程序、人工智慧執行期160或其他邊緣應用。執行期間，稽核記錄模組170可記錄五維度規封包 M5、治理判決 Vr、政策版本、節點識別、時間戳、調度結果、回滾狀態及其他事件資訊，並形成稽核日誌檔200。該稽核日誌檔200可進一步包含封包雜湊、父雜湊、啟動工作階段識別碼、回滾參照及重放防護參數。

    當應用程式執行完成、異常中止或接獲終止指令時，回復模組180即關閉相關服務、清除瞬態輸入輸出區130、移除執行期間注入之暫時性設定，並使系統狀態210回復至預定之初始狀態。若回復模組180判定執行事件具有重放風險、權限不連續、資料敏感性不符或拓樸不一致，則該執行事件被封存於隔離狀態，並禁止直接重放或恢復執行，直至取得預定之授權或稽核確認。
    """
).strip()


CLAIMS = [
    "一種基於唯讀啟動媒體與五維度規封包之邊緣運算執行環境注入、瞬態 I/O 記憶體調度及可稽核狀態回滾方法，其包含：以唯讀啟動媒體啟動邊緣運算裝置；擷取至少一環境參數；依據該環境參數及待執行動作產生一五維度規封包；依據該五維度規封包產生一治理判決；依據該治理判決注入一執行環境；於揮發性記憶體中建立一瞬態輸入輸出區；於該瞬態輸入輸出區中執行暫存讀寫；記錄一執行稽核資料；以及於終止條件成立時清除該瞬態輸入輸出區並回復執行狀態。",
    "如申請專利範圍第1項所述之方法，其中該五維度規封包包含節點或身份、資料敏感性、動作意圖、權限時間窗口及可逆性或公共價值。",
    "如申請專利範圍第1項所述之方法，其中該治理判決包含允許、稽核、警示、阻斷、要求人類確認及要求回滾之至少一者。",
    "如申請專利範圍第1項所述之方法，其中該五維度規封包用以決定該瞬態輸入輸出區之大小、生命週期、讀寫權限、清除條件及稽核等級。",
    "如申請專利範圍第1項所述之方法，其中該治理判決用以決定邊緣裝置執行、本地人工智慧模型執行、雲端轉送、延後執行、阻斷執行或隔離執行之其中一者。",
    "如申請專利範圍第1項所述之方法，其中該瞬態 I/O 記憶體調度係依據記憶體壓力、讀取壓力、寫入壓力、快取命中率、模型載入成本、網路成本、磁碟成本、資料敏感性及重放風險之至少一者執行。",
    "如申請專利範圍第1項所述之方法，其中該執行稽核資料至少包含啟動工作階段識別碼、五維度規封包雜湊值、父雜湊值、治理判決、政策版本、時間戳、回滾參照及重放防護參數。",
    "如申請專利範圍第1項所述之方法，其中當該待執行動作被判定為重放風險、權限不連續、資料敏感性不符或拓樸不一致時，將該待執行動作或其執行結果移入隔離狀態而非直接恢復執行。",
    "如申請專利範圍第1項所述之方法，其中該執行環境注入係將該環境參數或該五維度規封包所對應之參數注入容器協調檔案所定義之至少一服務。",
    "如申請專利範圍第1項所述之方法，其中該執行環境注入不將明文個人資料、金鑰、權杖、服務帳戶私密資料或其他私密認證資料寫入該唯讀啟動媒體或外部人工智慧推理內容。",
    "如申請專利範圍第1項所述之方法，其中該回復執行狀態包含停止服務程序、關閉容器、卸載或清除該瞬態輸入輸出區、還原環境參數、釋放資源及更新稽核紀錄。",
    "如申請專利範圍第1項所述之方法，其中該唯讀啟動媒體包含政策基準，且該政策基準用以約束該治理判決之產生。",
    "如申請專利範圍第1項所述之方法，其中該五維度規封包進一步包含政策雜湊值、節點簽章、封包有效期限或稽核識別碼。",
    "一種邊緣運算執行環境治理系統，其包含：唯讀啟動媒體、環境參數擷取模組、封包產生模組、AI 匝道判決模組、邊緣 I/O 記憶體調度器、人工智慧執行期、稽核記錄模組及回復模組；其中該封包產生模組產生五維度規封包，該 AI 匝道判決模組依據該五維度規封包產生治理判決，該邊緣 I/O 記憶體調度器依據該治理判決控制瞬態輸入輸出區之讀寫與生命週期。",
    "如申請專利範圍第14項所述之系統，其中該稽核記錄模組以雜湊鏈方式記錄執行事件，以供後續回滾、重放防護或稽核驗證。",
]


SYMBOLS = [
    ("100", "系統"),
    ("110", "輸入介面"),
    ("120", "封包產生模組"),
    ("130", "政策基準"),
    ("140", "AI 匝道判決模組"),
    ("150", "邊緣 I/O 記憶體調度器"),
    ("160", "AI runtime"),
    ("170", "稽核記錄模組"),
    ("180", "回復模組"),
    ("190", "邊緣節點或儲存節點"),
    ("M5", "五維度規封包"),
    ("Vr", "治理判決"),
    ("N", "節點或身份"),
    ("D", "資料敏感性"),
    ("A", "動作意圖"),
    ("W", "權限時間窗口"),
    ("R", "可逆性或公共價值"),
]


def ensure_out() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    (OUT / "figures").mkdir(exist_ok=True)


def write_text(path: Path, text: str) -> None:
    path.write_text(text.strip() + "\n", encoding="utf-8")


def svg_box(x, y, w, h, label, fill="#f8fafc", stroke="#334155"):
    return f'<rect x="{x}" y="{y}" width="{w}" height="{h}" rx="10" fill="{fill}" stroke="{stroke}" stroke-width="2"/><text x="{x+w/2}" y="{y+30}" text-anchor="middle" font-size="18" font-family="Noto Sans CJK TC, Arial" font-weight="700">{label}</text>'


def write_figures() -> None:
    fig1 = f"""<svg xmlns="http://www.w3.org/2000/svg" width="1400" height="760" viewBox="0 0 1400 760">
<style>text{{font-family:'Noto Sans CJK TC',Arial,sans-serif}} .small{{font-size:15px;font-weight:400}}</style>
<text x="700" y="42" text-anchor="middle" font-size="28" font-weight="700">圖1 系統架構方塊圖</text>
{svg_box(40,90,170,170,'110 輸入介面','#dcfce7')}
{svg_box(290,90,200,170,'120 封包產生模組','#fef3c7')}
{svg_box(570,90,230,170,'140 AI 匝道判決模組','#dbeafe')}
{svg_box(880,90,230,170,'150 邊緣 I/O 調度器','#dcfce7')}
{svg_box(1180,90,180,170,'160 AI runtime','#dbeafe')}
{svg_box(570,345,230,135,'170 稽核記錄模組','#fee2e2')}
{svg_box(880,345,230,135,'180 回復模組','#f1f5f9')}
{svg_box(1180,345,180,135,'190 邊緣/儲存節點','#fef3c7')}
{svg_box(570,570,230,100,'130 政策基準','#ede9fe')}
<defs><marker id="a" markerWidth="10" markerHeight="10" refX="8" refY="3" orient="auto"><path d="M0,0 L0,6 L9,3 z" fill="#111827"/></marker></defs>
<path d="M210 175 H290" stroke="#111827" stroke-width="3" marker-end="url(#a)"/><text x="248" y="160" font-size="16">資料</text>
<path d="M490 175 H570" stroke="#111827" stroke-width="3" marker-end="url(#a)"/><text x="525" y="160" font-size="16">M5</text>
<path d="M800 175 H880" stroke="#111827" stroke-width="3" marker-end="url(#a)"/><text x="835" y="160" font-size="16">Vr</text>
<path d="M1110 145 H1180" stroke="#111827" stroke-width="3" marker-end="url(#a)"/>
<path d="M1110 205 C1160 205 1160 410 1180 410" stroke="#111827" stroke-width="3" fill="none" marker-end="url(#a)"/>
<path d="M685 260 V345" stroke="#111827" stroke-width="3" marker-end="url(#a)"/>
<path d="M995 260 V345" stroke="#111827" stroke-width="3" marker-end="url(#a)"/>
<path d="M685 570 V480" stroke="#111827" stroke-width="3" marker-end="url(#a)"/>
<text x="125" y="150" class="small" text-anchor="middle">待執行動作</text><text x="125" y="175" class="small" text-anchor="middle">身份/權限資訊</text><text x="125" y="200" class="small" text-anchor="middle">節點狀態</text>
<text x="390" y="150" class="small" text-anchor="middle">產生 M5</text><text x="390" y="175" class="small" text-anchor="middle">M5=&lt;N,D,A,W,R&gt;</text>
<text x="685" y="150" class="small" text-anchor="middle">依 M5 與政策</text><text x="685" y="175" class="small" text-anchor="middle">輸出治理判決 Vr</text>
<text x="995" y="150" class="small" text-anchor="middle">控制排載、讀寫</text><text x="995" y="175" class="small" text-anchor="middle">重組、稽核、回滾</text>
</svg>"""
    (OUT / "figures" / "fig1_system_architecture.svg").write_text(fig1, encoding="utf-8")

    fig2 = """<svg xmlns="http://www.w3.org/2000/svg" width="1200" height="620" viewBox="0 0 1200 620">
<style>text{font-family:'Noto Sans CJK TC',Arial,sans-serif}.h{font-size:24px;font-weight:700}.s{font-size:16px}</style>
<text x="600" y="42" text-anchor="middle" class="h">圖2 五維度規向量封包資料結構示意圖</text>
<rect x="430" y="75" width="340" height="62" rx="12" fill="#f8fafc" stroke="#111827" stroke-width="2"/>
<text x="600" y="114" text-anchor="middle" font-size="26" font-weight="700">M5 = &lt; N, D, A, W, R &gt;</text>
<g transform="translate(55,190)">
<rect width="190" height="290" rx="12" fill="#dcfce7" stroke="#166534" stroke-width="2"/><text x="95" y="42" text-anchor="middle" font-size="22" font-weight="700">N</text><text x="95" y="75" text-anchor="middle" class="s">節點或身份</text><text x="95" y="122" text-anchor="middle" class="s">裝置</text><text x="95" y="152" text-anchor="middle" class="s">使用者</text><text x="95" y="182" text-anchor="middle" class="s">服務帳戶</text><text x="95" y="212" text-anchor="middle" class="s">容器</text><text x="95" y="242" text-anchor="middle" class="s">節點</text></g>
<g transform="translate(280,190)"><rect width="190" height="290" rx="12" fill="#fef3c7" stroke="#92400e" stroke-width="2"/><text x="95" y="42" text-anchor="middle" font-size="22" font-weight="700">D</text><text x="95" y="75" text-anchor="middle" class="s">資料敏感性</text><text x="95" y="122" text-anchor="middle" class="s">公開</text><text x="95" y="152" text-anchor="middle" class="s">內部</text><text x="95" y="182" text-anchor="middle" class="s">敏感</text><text x="95" y="212" text-anchor="middle" class="s">機密</text></g>
<g transform="translate(505,190)"><rect width="190" height="290" rx="12" fill="#dbeafe" stroke="#1e40af" stroke-width="2"/><text x="95" y="42" text-anchor="middle" font-size="22" font-weight="700">A</text><text x="95" y="75" text-anchor="middle" class="s">動作意圖</text><text x="95" y="122" text-anchor="middle" class="s">讀取</text><text x="95" y="152" text-anchor="middle" class="s">寫入</text><text x="95" y="182" text-anchor="middle" class="s">推論</text><text x="95" y="212" text-anchor="middle" class="s">封存</text><text x="95" y="242" text-anchor="middle" class="s">回滾</text></g>
<g transform="translate(730,190)"><rect width="190" height="290" rx="12" fill="#ede9fe" stroke="#6d28d9" stroke-width="2"/><text x="95" y="42" text-anchor="middle" font-size="22" font-weight="700">W</text><text x="95" y="75" text-anchor="middle" class="s">權限時間窗口</text><text x="95" y="122" text-anchor="middle" class="s">起始/結束</text><text x="95" y="152" text-anchor="middle" class="s">授權場域</text><text x="95" y="182" text-anchor="middle" class="s">有效期限</text></g>
<g transform="translate(955,190)"><rect width="190" height="290" rx="12" fill="#ffedd5" stroke="#c2410c" stroke-width="2"/><text x="95" y="42" text-anchor="middle" font-size="22" font-weight="700">R</text><text x="95" y="75" text-anchor="middle" class="s">可逆性或公共價值</text><text x="95" y="122" text-anchor="middle" class="s">可逆/不可逆</text><text x="95" y="152" text-anchor="middle" class="s">回復成本</text><text x="95" y="182" text-anchor="middle" class="s">人類確認</text><text x="95" y="212" text-anchor="middle" class="s">公共風險</text></g>
</svg>"""
    (OUT / "figures" / "fig2_m5_packet.svg").write_text(fig2, encoding="utf-8")

    for idx, title in [
        (3, "AI 匝道判決流程圖"),
        (4, "邊緣 I/O 記憶體調度流程圖"),
        (5, "稽核紀錄與回滾流程示意圖"),
    ]:
        steps = {
            3: ["接收待執行動作", "產生 M5", "讀取政策基準", "計算治理判決 Vr", "輸出 allow/audit/warn/block"],
            4: ["接收治理判決 Vr", "判斷排載記憶體區段", "判斷讀取上下文片段", "判斷重組上下文", "建立稽核紀錄", "必要時觸發回復"],
            5: ["讀取稽核紀錄", "確認回滾條件", "撤銷記憶體掛載", "刪除暫存上下文", "封存輸出結果", "恢復至先前狀態", "更新稽核紀錄"],
        }[idx]
        y = 90
        rects = []
        arrows = []
        for i, step in enumerate(steps, 1):
            rects.append(f'<rect x="360" y="{y}" width="480" height="56" rx="12" fill="#f8fafc" stroke="#334155" stroke-width="2"/><text x="600" y="{y+36}" text-anchor="middle" font-size="20">S{idx}{i:02d} {step}</text>')
            if i < len(steps):
                arrows.append(f'<path d="M600 {y+56} V{y+88}" stroke="#111827" stroke-width="3" marker-end="url(#a)"/>')
            y += 95
        svg = f"""<svg xmlns="http://www.w3.org/2000/svg" width="1200" height="{max(700, y+40)}" viewBox="0 0 1200 {max(700, y+40)}">
<defs><marker id="a" markerWidth="10" markerHeight="10" refX="8" refY="3" orient="auto"><path d="M0,0 L0,6 L9,3 z" fill="#111827"/></marker></defs>
<style>text{{font-family:'Noto Sans CJK TC',Arial,sans-serif}}</style>
<text x="600" y="42" text-anchor="middle" font-size="28" font-weight="700">圖{idx} {title}</text>
{''.join(rects)}
{''.join(arrows)}
</svg>"""
        (OUT / "figures" / f"fig{idx}.svg").write_text(svg, encoding="utf-8")


def application_md(kind: str) -> str:
    if kind == "individual":
        applicant = f"{INVENTOR['name']}（自然人）"
        applicant_id = INVENTOR["id"]
        applicant_address = INVENTOR["address"]
    else:
        applicant = f"{BUSINESS['name']}（商號）"
        applicant_id = BUSINESS["ubn"]
        applicant_address = BUSINESS["address"]
    return dedent(
        f"""
        # 發明專利申請書草稿（{kind} 版本）

        > 本檔供網路送件前轉填官方申請書或交專利師確認。實際欄位仍以智慧財產局電子申請系統及最新表單為準。

        ## 申請案資料

        - 申請種類：發明專利
        - 發明名稱：{TITLE}
        - 指定代表圖：圖1
        - 代表圖符號簡單說明：100 系統；110 輸入介面；120 封包產生模組；130 政策基準；140 AI 匝道判決模組；150 邊緣 I/O 記憶體調度器；160 AI runtime；170 稽核記錄模組；180 回復模組；190 邊緣節點或儲存節點。

        ## 申請人草稿

        - 申請人：{applicant}
        - 識別資料：{applicant_id}
        - 地址：{applicant_address}

        ## 發明人

        - 發明人：{INVENTOR['name']}
        - 識別資料：{INVENTOR['id']}
        - 地址：{INVENTOR['address']}

        ## 商號/實施場域備註

        - 商號：{BUSINESS['name']}
        - 統一編號：{BUSINESS['ubn']}
        - 地址：{BUSINESS['address']}
        - 負責人：{BUSINESS['representative']}
        - 備註：本發明係於邊緣運算、POS/物業/社區治理系統開發過程中衍生，商號資訊可作為實施場域或申請人版本之一，實際申請人歸屬請於送件前確認。

        ## 檢附文件

        - 摘要
        - 說明書
        - 申請專利範圍
        - 圖式
        - 符號說明
        """
    ).strip()


def write_docs() -> None:
    write_text(
        OUT / "README_送件包.md",
        f"""
        # {SHORT_TITLE} 送件草稿包

        產生時間：{NOW.isoformat()}

        本資料包用途：供智慧財產局網路送件前整理、轉填官方申請書，或交由專利師進行正式送件前檢核。此資料包不是法律意見，也不保證核准。

        ## 主要檔案

        - `01_發明專利申請書草稿_自然人申請.md`
        - `01B_發明專利申請書草稿_商號申請.md`
        - `02_摘要.md`
        - `03_說明書.md`
        - `04_申請專利範圍.md`
        - `05_圖式簡單說明.md`
        - `06_符號說明.md`
        - `07_發明揭露與營業秘密邊界備忘錄.md`
        - `08_送件檢核表.md`
        - `09_合併送件草稿.docx`
        - `figures/*.svg`

        ## 送件提醒

        1. 實際申請書請以智慧財產局電子申請系統及最新官方表單為準。
        2. 申請人可選自然人或商號版本，請送件前確認權利歸屬。
        3. 本檔已避免寫入金鑰、token、service account JSON 或系統秘密。
        4. 若要正式送件，建議由專利師再做前案檢索、請求項修飾與圖式格式檢查。
        """,
    )
    write_text(OUT / "01_發明專利申請書草稿_自然人申請.md", application_md("individual"))
    write_text(OUT / "01B_發明專利申請書草稿_商號申請.md", application_md("business"))
    write_text(OUT / "02_摘要.md", "# 摘要\n\n" + ABSTRACT + "\n\n## 指定代表圖\n圖1。")
    write_text(OUT / "03_說明書.md", SPECIFICATION)
    write_text(OUT / "04_申請專利範圍.md", "# 申請專利範圍\n\n" + "\n\n".join(f"{i}. {c}" for i, c in enumerate(CLAIMS, 1)))
    write_text(
        OUT / "05_圖式簡單說明.md",
        "# 圖式簡單說明\n\n"
        "圖1係本發明之系統架構方塊圖。\n\n"
        "圖2係本發明之五維度規向量封包資料結構示意圖。\n\n"
        "圖3係本發明之 AI 匝道判決流程圖。\n\n"
        "圖4係本發明之邊緣 I/O 記憶體調度流程圖。\n\n"
        "圖5係本發明之稽核紀錄與回滾流程示意圖。\n",
    )
    write_text(OUT / "06_符號說明.md", "# 符號說明\n\n" + "\n".join(f"- {a}：{b}" for a, b in SYMBOLS))
    write_text(
        OUT / "07_發明揭露與營業秘密邊界備忘錄.md",
        f"""
        # 發明揭露與營業秘密邊界備忘錄

        ## 可揭露技術核心

        - 唯讀啟動媒體啟動邊緣節點。
        - 五維度規封包 M5 = <N,D,A,W,R>。
        - 治理判決 Vr = allow/audit/warn/block/require_human_confirmation/rollback_required。
        - 瞬態 I/O 記憶體調度。
        - 稽核雜湊鏈、回滾參照、重放防護。
        - 無明文上下文輸出至外部 AI。

        ## 建議保留為營業秘密

        - 實際部署節點清單、IP、金鑰與憑證。
        - 會員或客戶明文資料。
        - 未公開之模型前綴、權重、策略閾值、商業定價。
        - 實際 Odoo 資料庫內容與服務帳戶設定。

        ## 權利歸屬待確認

        - 申請人採自然人或商號。
        - 商號作為實施場域或權利人之角色。
        - 若未來納入協會、公司、合作開發者，應另行簽署權利讓與或授權文件。
        """,
    )
    write_text(
        OUT / "08_送件檢核表.md",
        """
        # 送件檢核表

        - [ ] 申請人版本確認：自然人 / 商號 / 其他法人。
        - [ ] 發明人姓名、地址、識別資料確認。
        - [ ] 摘要字數與代表圖確認。
        - [ ] 說明書章節順序確認。
        - [ ] 申請專利範圍請求項檢查。
        - [ ] 圖式圖號、符號、說明一致。
        - [ ] 圖式內文字是否符合送件規範，由專利師或送件前再確認。
        - [ ] 前案檢索完成。
        - [ ] 未公開營業秘密已移除。
        - [ ] 無金鑰、token、密碼、service account JSON。
        - [ ] 電子申請表單依最新官方格式轉填。
        """,
    )


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_paragraphs(doc, text):
    for block in text.split("\n\n"):
        block = block.strip()
        if not block:
            continue
        if block.startswith("【") and block.endswith("】"):
            doc.add_heading(block.strip("【】"), level=1)
        else:
            for line in block.splitlines():
                if line.strip():
                    doc.add_paragraph(line.strip())


def write_docx() -> None:
    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(12)
    doc.add_heading(TITLE, level=0)
    doc.add_paragraph(f"送件草稿版本：{NOW.strftime('%Y-%m-%d %H:%M:%S %z')}")
    doc.add_paragraph("注意：本檔為專利送件前草稿，實際送件仍應依智慧財產局電子申請系統及最新表單轉填。")
    doc.add_page_break()
    doc.add_heading("摘要", level=1)
    doc.add_paragraph(ABSTRACT)
    doc.add_paragraph("指定代表圖：圖1。")
    doc.add_page_break()
    add_paragraphs(doc, SPECIFICATION)
    doc.add_page_break()
    doc.add_heading("申請專利範圍", level=1)
    for i, claim in enumerate(CLAIMS, 1):
        doc.add_paragraph(f"{i}. {claim}")
    doc.add_page_break()
    doc.add_heading("符號說明", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "符號"
    table.rows[0].cells[1].text = "說明"
    for sym, desc in SYMBOLS:
        row = table.add_row().cells
        row[0].text = sym
        row[1].text = desc
    doc.add_page_break()
    doc.add_heading("圖式簡單說明", level=1)
    for line in (OUT / "05_圖式簡單說明.md").read_text(encoding="utf-8").splitlines()[2:]:
        if line.strip():
            doc.add_paragraph(line.strip())
    doc.save(OUT / "09_合併送件草稿.docx")


def write_manifest() -> None:
    files = []
    for p in sorted(OUT.rglob("*")):
        if p.is_file():
            h = hashlib.sha256(p.read_bytes()).hexdigest()
            files.append({"path": str(p.relative_to(OUT)), "sha256": h, "bytes": p.stat().st_size})
    write_text(OUT / "MANIFEST.json", json.dumps({"generated_at": NOW.isoformat(), "title": TITLE, "files": files}, ensure_ascii=False, indent=2))


def main() -> None:
    ensure_out()
    write_figures()
    write_docs()
    write_docx()
    write_manifest()
    print(OUT)


if __name__ == "__main__":
    main()
